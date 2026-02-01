
import re
from docx import Document
from docx.shared import Pt, Inches

def parse_markdown_to_docx(md_file, docx_file):
    doc = Document()
    
    # Set default style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    table_mode = False
    current_table = None
    
    for line in lines:
        stripped_line = line.strip()
        
        # Skip empty lines if not in table mode (tables need checking)
        if not stripped_line:
            if table_mode:
                table_mode = False
                current_table = None
            continue

        # Tables
        if stripped_line.startswith('|'):
            if not table_mode:
                # Start of a new table
                table_mode = True
                # Count columns
                headers = [h.strip() for h in stripped_line.split('|') if h.strip()]
                current_table = doc.add_table(rows=1, cols=len(headers))
                current_table.style = 'Table Grid'
                hdr_cells = current_table.rows[0].cells
                for i, header in enumerate(headers):
                    hdr_cells[i].text = header
                    # Make bold
                    for paragraph in hdr_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
            else:
                # Existing table row
                # Check if it's a separator line
                is_separator = True
                parts = [p.strip() for p in stripped_line.split('|') if p.strip()]
                for part in parts:
                    if not re.match(r'^[\-\:]+$', part):
                        is_separator = False
                        break
                
                if is_separator:
                    continue
                
                # Add data row
                row_cells = current_table.add_row().cells
                # Handle cases where row might have different col count due to escaped pipes (unlikely here) or empty end pipes
                # Split logic: split by pipe, ignore first/last if empty
                raw_cells = stripped_line.split('|')
                # Filter out the empty strings often found at start/end of |...|...|
                if raw_cells[0].strip() == '': raw_cells.pop(0)
                if raw_cells and raw_cells[-1].strip() == '': raw_cells.pop(-1)
                
                for i, cell_text in enumerate(raw_cells):
                    if i < len(row_cells):
                        row_cells[i].text = cell_text.strip()
            continue
        else:
            if table_mode:
                table_mode = False
                current_table = None

        # Headers
        if stripped_line.startswith('# '):
            doc.add_heading(stripped_line[2:], level=1)
        elif stripped_line.startswith('## '):
            doc.add_heading(stripped_line[3:], level=2)
        elif stripped_line.startswith('### '):
            doc.add_heading(stripped_line[4:], level=3)
        elif stripped_line.startswith('#### '):
            doc.add_heading(stripped_line[5:], level=4)
        
        # Bullet Points
        elif stripped_line.startswith('* '):
            p = doc.add_paragraph(stripped_line[2:], style='List Bullet')
            apply_formatting(p)
        elif stripped_line.startswith('- '):
            p = doc.add_paragraph(stripped_line[2:], style='List Bullet')
            apply_formatting(p)
        else:
            # Regular Paragraph
            # Handle blockquotes
            if stripped_line.startswith('>'):
                p = doc.add_paragraph(stripped_line.replace('>', '').strip())
                p.italic = True 
            else:
                p = doc.add_paragraph(stripped_line)
                apply_formatting(p)

    doc.save(docx_file)
    print(f"Successfully converted {md_file} to {docx_file}")

def apply_formatting(paragraph):
    # Handle bold **text**
    # This is a basic parser approach: split by ** and toggle bold
    # It assumes even number of ** markers
    text = paragraph.text
    if '**' not in text and '$$' not in text:
        return
        
    # Clear existing runs to rebuild with formatting
    paragraph.clear()
    
    # Very simple state machine for bold
    parts = re.split(r'(\*\*|$$)', text) # Split by markers
    # This might be too complex for regex split; let's do manual scan or just simple Bold handling
    # Let's try a regex for **content**
    
    # Re-impl: iterate and find matches
    # Find all bold chunks
    
    # Basic strategy: Alternating plain/bold/plain
    # Only if structure is simple like "Text **Bold** Text"
    
    # Let's traverse the text
    cursor = 0
    is_bold = False
    
    # We will rebuild the paragraph runs
    # Note: re.split with capturing group keeps the delimiters
    tokens = re.split(r'(\*\*)', text)
    
    for token in tokens:
        if token == '**':
            is_bold = not is_bold
        else:
            if token:
                run = paragraph.add_run(token)
                run.font.bold = is_bold

if __name__ == "__main__":
    parse_markdown_to_docx('cost_analysis.md', 'cost_analysis.docx')
